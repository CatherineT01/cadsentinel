# spec_parser.py

"""
cadsentinel.etl.spec_parser
----------------------------
Document parsing and chunking for spec ingestion.
Supports PDF, DOCX, and plain text (.txt) files.

Chunking strategy: fixed-size with overlap.
  - Chunk size: ~2000 characters
  - Overlap:    200 characters (preserves context at boundaries)

This module has no database or LLM logic — pure I/O and text splitting.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from .spec_models import SpecChunk, SpecDocument

log = logging.getLogger(__name__)

CHUNK_SIZE    = 2000   # characters per chunk
CHUNK_OVERLAP = 200    # overlap between consecutive chunks
MIN_CHUNK_LEN = 50     # discard chunks shorter than this


class ParseError(RuntimeError):
    """Raised when a document cannot be parsed."""


# ── Public entry point ───────────────────────────────────────── #

def parse_document(file_path: str | Path) -> tuple[SpecDocument, list[SpecChunk]]:
    """
    Parse a specification document into a SpecDocument + list of SpecChunks.

    Args:
        file_path: Path to the PDF, DOCX, or TXT file.

    Returns:
        (SpecDocument, [SpecChunk, ...])

    Raises:
        ParseError: If the file cannot be read or format is unsupported.
    """
    path = Path(file_path).resolve()
    if not path.exists():
        raise ParseError(f"File not found: {path}")

    suffix = path.suffix.lower()

    if suffix == ".pdf":
        raw_text, page_map = _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        raw_text, page_map = _parse_docx(path)
    elif suffix == ".txt":
        raw_text, page_map = _parse_txt(path)
    else:
        raise ParseError(
            f"Unsupported file format: '{suffix}'. "
            "Supported formats: .pdf, .docx, .txt"
        )

    if not raw_text.strip():
        raise ParseError(f"Document appears to be empty: {path}")

    doc = SpecDocument(
        filename    = path.name,
        source_type = suffix.lstrip("."),
        raw_text    = raw_text,
        title       = _infer_title(raw_text, path.name),
    )

    chunks = _chunk_text(raw_text, page_map)

    log.info(
        "Parsed '%s': %d chars, %d chunks",
        path.name, len(raw_text), len(chunks)
    )

    return doc, chunks


# ── Format-specific parsers ──────────────────────────────────── #

def _parse_pdf(path: Path) -> tuple[str, dict[int, int]]:
    """
    Extract text from a PDF using pymupdf (fitz).
    Returns (full_text, page_map) where page_map maps char_offset -> page_number.
    """
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise ParseError(
            "pymupdf is required for PDF parsing. Install it: pip install pymupdf"
        ) from exc

    try:
        doc = fitz.open(str(path))
    except Exception as exc:
        raise ParseError(f"Failed to open PDF '{path}': {exc}") from exc

    pages: list[str] = []
    page_map: dict[int, int] = {}   # char_offset -> page_number
    char_offset = 0

    for page_num, page in enumerate(doc, start=1):
        try:
            text = page.get_text("text") or ""
        except Exception:
            log.warning("Failed to extract text from page %d of %s", page_num, path.name)
            text = ""

        page_map[char_offset] = page_num
        pages.append(text)
        char_offset += len(text)

    doc.close()
    return "\n".join(pages), page_map


def _parse_docx(path: Path) -> tuple[str, dict[int, int]]:
    """
    Extract text from a DOCX file using python-docx.
    DOCX has no page concept — page_map will be empty.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ParseError(
            "python-docx is required for DOCX parsing. Install it: pip install python-docx"
        ) from exc

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ParseError(f"Failed to open DOCX '{path}': {exc}") from exc

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    raw_text = "\n".join(paragraphs)
    return raw_text, {}


def _parse_txt(path: Path) -> tuple[str, dict[int, int]]:
    """Read a plain text file. Try UTF-8, fall back to latin-1."""
    try:
        raw_text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            raw_text = path.read_text(encoding="latin-1")
        except Exception as exc:
            raise ParseError(f"Failed to read text file '{path}': {exc}") from exc
    except Exception as exc:
        raise ParseError(f"Failed to read text file '{path}': {exc}") from exc

    return raw_text, {}


# ── Chunking ─────────────────────────────────────────────────── #

def _chunk_text(
    text: str,
    page_map: dict[int, int],
    chunk_size:    int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[SpecChunk]:
    chunks: list[SpecChunk] = []
    start = 0
    index = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        # Try to break at a paragraph boundary within the last 20% of the chunk
        if end < len(text):
            search_from = start + int(chunk_size * 0.8)
            newline_pos = text.rfind("\n\n", search_from, end)
            if newline_pos != -1:
                end = newline_pos + 2
            else:
                newline_pos = text.rfind("\n", search_from, end)
                if newline_pos != -1:
                    end = newline_pos + 1

        chunk_text = text[start:end].strip()

        if len(chunk_text) >= MIN_CHUNK_LEN:
            page_num     = _page_for_offset(start, page_map)
            section_hint = _detect_heading(chunk_text)
            chunks.append(SpecChunk(
                chunk_index  = index,
                raw_text     = chunk_text,
                char_start   = start,
                char_end     = end,
                source_page  = page_num,
                section_hint = section_hint,
            ))
            index += 1

        # Advance start — guarantee forward progress
        next_start = end - chunk_overlap
        if next_start <= start:
            next_start = start + chunk_size
        start = next_start

    return chunks


def _page_for_offset(char_offset: int, page_map: dict[int, int]) -> Optional[int]:
    """Return the page number for a given character offset."""
    if not page_map:
        return None
    best_page = None
    for offset, page in page_map.items():
        if offset <= char_offset:
            best_page = page
        else:
            break
    return best_page


def _detect_heading(text: str) -> Optional[str]:
    """
    Heuristic: if the first non-empty line of a chunk is short and
    doesn't end with a period, it's probably a section heading.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    if not lines:
        return None
    first_line = lines[0]
    if len(first_line) < 120 and not first_line.endswith("."):
        return first_line
    return None


def _infer_title(text: str, filename: str) -> str:
    """
    Try to infer a document title from the first meaningful line of text.
    Falls back to the filename stem.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:10]:
        if 5 < len(line) < 150 and not line.endswith("."):
            return line
    return Path(filename).stem.replace("_", " ").replace("-", " ").title()